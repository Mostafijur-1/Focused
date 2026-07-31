from __future__ import annotations

import argparse
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(106, 118, 130)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build the Focused design-system Pandoc reference DOCX.")
    parser.add_argument("base", type=Path)
    parser.add_argument("output", type=Path)
    return parser.parse_args()


def set_style_font(style, *, name: str, size: float, color: RGBColor | None = None, bold: bool | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), "Nirmala UI")
    rfonts.set(qn("w:cs"), "Nirmala UI")


def set_spacing(style, *, before: float, after: float, line: float) -> None:
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.widow_control = True


def set_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def patch_numbering(doc: Document) -> None:
    numbering = doc.part.numbering_part.element
    for level in numbering.xpath(".//w:lvl[@w:ilvl='0']"):
        ppr = level.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            level.append(ppr)
        ind = ppr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            ppr.append(ind)
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        spacing = ppr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            ppr.append(spacing)
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")


def main() -> None:
    args = parse_args()
    args.output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(args.base, args.output)
    doc = Document(args.output)

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run("FOCUSED | BANGLA-FIRST UI/UX DESIGN SYSTEM")
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED

        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.clear()
        set_page_number(footer_paragraph)

    styles = doc.styles
    set_style_font(styles["Normal"], name="Calibri", size=11, color=RGBColor(32, 37, 43))
    set_spacing(styles["Normal"], before=0, after=6, line=1.25)

    style_specs = {
        "Title": (30, NAVY, 0, 8, 1.10, True),
        "Subtitle": (15, MUTED, 0, 18, 1.20, False),
        "Heading 1": (16, BLUE, 18, 10, 1.10, True),
        "Heading 2": (13, BLUE, 14, 7, 1.15, True),
        "Heading 3": (12, DARK_BLUE, 10, 5, 1.20, True),
        "Caption": (9, MUTED, 4, 8, 1.15, False),
    }
    for name, (size, color, before, after, line, bold) in style_specs.items():
        if name not in styles:
            continue
        style = styles[name]
        set_style_font(style, name="Calibri", size=size, color=color, bold=bold)
        set_spacing(style, before=before, after=after, line=line)
        if name.startswith("Heading"):
            style.paragraph_format.keep_with_next = True

    for name in ("Body Text", "First Paragraph", "Block Text"):
        if name in styles:
            set_style_font(styles[name], name="Calibri", size=11, color=RGBColor(32, 37, 43))
            set_spacing(styles[name], before=0, after=6, line=1.25)

    for name in ("List Bullet", "List Number", "Compact"):
        if name in styles:
            set_style_font(styles[name], name="Calibri", size=11, color=RGBColor(32, 37, 43))
            set_spacing(styles[name], before=0, after=4, line=1.25)
            styles[name].paragraph_format.left_indent = Inches(0.375)
            styles[name].paragraph_format.first_line_indent = Inches(-0.188)

    patch_numbering(doc)
    doc.save(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
